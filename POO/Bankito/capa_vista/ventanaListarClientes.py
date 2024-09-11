from PySide6.QtWidgets import QApplication, QMainWindow,QDialog, QTableWidget, QTableWidgetItem, QVBoxLayout, QPushButton, QWidget
import sys
from capa_vista.VentanaListado_ui import Ui_VentanaListadoDeClientes
from capa_Negocio.ClienteDB import *
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

class VentanaListarClientes(QDialog,Ui_VentanaListadoDeClientes):
    def __init__(self):
        super(VentanaListarClientes, self).__init__()

        # Llamar al método setupUi para inicializar los widgets definidos en el archivo .ui
        self.setupUi(self)

        # Establecer título y dimensiones de la ventana      
        self.setWindowTitle("Listado de Clientes")
        self.setGeometry(100, 100, 600, 400)

        #traer los cliente desde el backend
        gestionClientes=ClientDB()
        self.clientes = gestionClientes.traer_Todos()


        # Crear el QTableWidget para mostrar los clientes
        self.tablaClientes.setRowCount(len(self.clientes))  # Cantidad de clientes
        self.tablaClientes.setColumnCount(3)  # Columnas: Id, Nombre, Domicilio
        self.tablaClientes.setHorizontalHeaderLabels(["Id", "Nombre", "Domicilio"])
       
        # Llenar la tabla con los datos de los clientes
        for row, cliente in enumerate(self.clientes):
            id_cliente = str(cliente[0])  # Asegúrate de convertir a cadena
            nombre_cliente = cliente[1]
            domicilio_cliente = cliente[2]
            self.tablaClientes.setItem(row, 0, QTableWidgetItem(id_cliente))
            self.tablaClientes.setItem(row, 1, QTableWidgetItem(nombre_cliente))
            self.tablaClientes.setItem(row, 2, QTableWidgetItem(domicilio_cliente))

        # Botón para exportar a Excel
        self.pushButton.setText("Crear Excel")
        self.pushButton.clicked.connect(self.exportar_excel)

        # Botón para exportar a Word
        self.pushButton_2.setText("Crear Word")
        self.pushButton_2.clicked.connect(self.exportar_word)
        

        # Botón para exportar a PDF
        self.pushButton_3.setText("Crear PDF")
        self.pushButton_3.clicked.connect(self.exportar_pdf)
    
    
    def exportar_excel(self):
        # Lógica para exportar los datos a Excel usando libreria pandas
        print("Exportando a Excel...")
        clientesTabla = []
        for row in range(self.tablaClientes.rowCount()):
            cliente = {
                "Id": self.tablaClientes.item(row, 0).text(),
                "Nombre": self.tablaClientes.item(row, 1).text(),
                "Domicilio": self.tablaClientes.item(row, 2).text(),
            }
            clientesTabla.append(cliente)
        
        df = pd.DataFrame(clientesTabla)  #crea un dataframe con los datos de la tabla
        df.to_excel("clientes.xlsx", index=False)
        print("Archivo Excel creado.")


    def exportar_word(self):
        # Lógica para exportar los datos a Word
        print("Exportando a WORD...")
        doc = Document()
        doc.add_heading('Listado de Clientes', 0)

        for row in range(self.tablaClientes.rowCount()):
            id_cliente = self.tablaClientes.item(row, 0).text()
            nombre_cliente = self.tablaClientes.item(row, 1).text()
            domicilio_cliente = self.tablaClientes.item(row, 2).text()
            doc.add_paragraph(f"{id_cliente} - {nombre_cliente} - {domicilio_cliente}")
        
        doc.save("clientes.docx")
        print("Archivo Word creado.")


    def exportar_pdf(self):
        # Lógica para exportar los datos a PDF
        print("Exportando a PDF...")

        c = canvas.Canvas("clientes.pdf", pagesize=letter)
        c.drawString(100, 750, "Listado de Clientes")

        y = 730
        for row in range(self.tablaClientes.rowCount()):
            id_cliente = self.tablaClientes.item(row, 0).text()
            nombre_cliente = self.tablaClientes.item(row, 1).text()
            domicilio_cliente = self.tablaClientes.item(row, 2).text()
            c.drawString(100, y, f"{id_cliente} - {nombre_cliente} - {domicilio_cliente}")
            y -= 20
        
        c.save()
        print("Archivo PDF creado.")

